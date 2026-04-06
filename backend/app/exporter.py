import json
import logging
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.config import EXPORTS_DIR
from app.db import list_evidence_events

logger = logging.getLogger(__name__)


def _ensure_dir() -> None:
    Path(EXPORTS_DIR).mkdir(parents=True, exist_ok=True)


def _safe_filename(value: str) -> str:
    cleaned = ''.join(ch for ch in value if ch not in r'\\/:*?"<>|').strip()
    return cleaned[:60] or 'export'


def _set_font(run, size_pt: int = 11, bold: bool = False) -> None:
    run.font.size = Pt(size_pt)
    run.font.bold = bold


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _normalize_result_payload(result: dict) -> tuple[str, dict[str, Any], list[str], str]:
    result_type = result.get('result_type', 'unknown')
    content = result.get('content_json', {})
    if isinstance(content, str):
        content = json.loads(content)
    evidence = result.get('evidence_refs', [])
    if isinstance(evidence, str):
        evidence = json.loads(evidence)
    title = content.get('title', result_type)
    return result_type, content, evidence, title


def _write_json_export(name: str, payload: dict[str, Any]) -> str:
    _ensure_dir()
    file_path = Path(EXPORTS_DIR) / f'{_safe_filename(name)}.json'
    file_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding='utf-8')
    return str(file_path)


def _add_kv_block(doc: Document, items: list[tuple[str, Any]]) -> None:
    for label, value in items:
        if value is None:
            continue
        doc.add_paragraph(f'{label}：{value}')


def _add_list_block(doc: Document, heading: str, items: list[Any]) -> None:
    if not items:
        return
    _add_heading(doc, heading, level=2)
    for item in items:
        doc.add_paragraph(str(item), style='List Bullet')


def _render_lesson_outline(doc: Document, content: dict[str, Any]) -> None:
    _add_list_block(doc, '教学目标', content.get('teaching_goals', []))
    _add_list_block(doc, '教学重点', content.get('key_points', []))
    _add_list_block(doc, '教学难点', content.get('difficult_points', []))
    activities = content.get('activity_flow', [])
    if activities:
        _add_heading(doc, '活动流程', level=2)
        for step in activities:
            if isinstance(step, dict):
                summary = f"步骤{step.get('step', '-')}: {step.get('name', '')}（{step.get('duration', '')}）"
                doc.add_paragraph(summary, style='List Number')
                if step.get('description'):
                    doc.add_paragraph(step['description'])
            else:
                doc.add_paragraph(str(step), style='List Number')


def _render_question_analysis(doc: Document, content: dict[str, Any]) -> None:
    _add_kv_block(doc, [('题目', content.get('question_text'))])
    _add_list_block(doc, '分析步骤', content.get('analysis_steps', []))
    hint_ladder = content.get('hint_ladder') or {}
    if hint_ladder:
        _add_heading(doc, '提示阶梯', level=2)
        for key in ['hint_1', 'hint_2', 'hint_3', 'answer']:
            if hint_ladder.get(key):
                doc.add_paragraph(f'{key}：{hint_ladder[key]}')


def _render_guided_explain(doc: Document, content: dict[str, Any]) -> None:
    _add_kv_block(
        doc,
        [
            ('Bloom 层级', content.get('bloom_level')),
            ('当前提示层', content.get('current_hint_level')),
            ('讲解内容', content.get('hint_content')),
            ('下一挑战提示', content.get('next_challenge_hint')),
        ],
    )


def export_result(result: dict, actor_role: str = 'teacher') -> dict[str, str | None]:
    if result.get('status') != 'ready':
        raise ValueError('only ready results can be exported')
    if result.get('review_status') == 'rejected':
        raise ValueError('rejected results cannot be exported')

    result_type, content, evidence, title = _normalize_result_payload(result)

    _ensure_dir()
    date_str = datetime.now(timezone.utc).strftime('%Y-%m-%d')

    try:
        doc = Document()
        cover = doc.add_paragraph()
        cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cover.add_run(title)
        _set_font(run, 16, bold=True)

        doc.add_paragraph(f'类型：{result_type}  |  日期：{date_str}  |  角色：{actor_role}')
        doc.add_paragraph('')

        if result_type == 'lesson_outline':
            _render_lesson_outline(doc, content)
        elif result_type == 'question_analysis':
            _render_question_analysis(doc, content)
        elif result_type == 'guided_explain':
            _render_guided_explain(doc, content)
        else:
            _add_heading(doc, '结果内容', level=2)
            doc.add_paragraph(json.dumps(content, ensure_ascii=False, indent=2))

        if evidence:
            _add_heading(doc, '证据链附录', level=2)
            for ref in evidence:
                doc.add_paragraph(f'• {ref}', style='List Bullet')

        file_name = _safe_filename(f'{title}_{result_type}_{date_str}') + '.docx'
        file_path = str(Path(EXPORTS_DIR) / file_name)
        doc.save(file_path)
        logger.info('exported result to %s', file_path)
        return {'format': 'docx', 'file_path': file_path, 'fallback_message': None}
    except Exception as exc:
        logger.exception('docx export failed, falling back to json')
        payload = {
            'result': result,
            'actor_role': actor_role,
            'exported_at': datetime.now(timezone.utc).isoformat(),
            'docx_error': str(exc),
        }
        file_path = _write_json_export(f'{title}_{result_type}_{date_str}', payload)
        return {
            'format': 'json',
            'file_path': file_path,
            'fallback_message': 'Word 导出失败，已提供 JSON 备份',
        }


def export_session_payload(session_id: str, result: dict | None, actor_role: str = 'teacher') -> dict[str, Any]:
    events = list_evidence_events(session_id)
    normalized_result = None
    if result:
        _, content, evidence, _ = _normalize_result_payload(result)
        normalized_result = {**result, 'content_json': content, 'evidence_refs': evidence}
    return {
        'session_id': session_id,
        'actor_role': actor_role,
        'exported_at': datetime.now(timezone.utc).isoformat(),
        'events': events,
        'result': normalized_result,
    }


def export_session(session_id: str, result: dict | None, actor_role: str = 'teacher') -> dict[str, str | None]:
    _ensure_dir()
    events = list_evidence_events(session_id)
    date_str = datetime.now(timezone.utc).strftime('%Y-%m-%d')

    try:
        doc = Document()
        cover = doc.add_paragraph()
        cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cover.add_run(f'会话记录 {session_id[:8]}')
        _set_font(run, 14, bold=True)
        doc.add_paragraph(f'日期：{date_str}  |  角色：{actor_role}')
        doc.add_paragraph('')

        _add_heading(doc, '证据事件时间线', level=2)
        for event in events:
            payload = event.get('payload', {})
            if isinstance(payload, str):
                payload = json.loads(payload)
            doc.add_paragraph(
                f'[{event["created_at"]}] {event["event_type"]} / {event["actor"]}  bloom={event.get("bloom_level", "-")}'
            )
            if payload:
                doc.add_paragraph(json.dumps(payload, ensure_ascii=False, indent=2))

        if result:
            _, content, _, _ = _normalize_result_payload(result)
            _add_heading(doc, '最终结果对象', level=2)
            doc.add_paragraph(json.dumps(content, ensure_ascii=False, indent=2))

        file_name = _safe_filename(f'session_{session_id[:8]}_{date_str}') + '.docx'
        file_path = str(Path(EXPORTS_DIR) / file_name)
        doc.save(file_path)
        logger.info('exported session to %s', file_path)
        return {'format': 'docx', 'file_path': file_path, 'fallback_message': None}
    except Exception as exc:
        logger.exception('session docx export failed, falling back to json')
        payload = export_session_payload(session_id, result, actor_role)
        payload['docx_error'] = str(exc)
        file_path = _write_json_export(f'session_{session_id[:8]}_{date_str}', payload)
        return {
            'format': 'json',
            'file_path': file_path,
            'fallback_message': 'Word 导出失败，已提供 JSON 备份',
        }
